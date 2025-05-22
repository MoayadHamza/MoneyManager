import tkinter as tk
from tkinter import ttk, messagebox
import tkinter.font as tkFont
from database import *


class MoneyManagerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Money Manager")
        initialize_db()  # Initialize the database
        
        # Input Fields
        self.amount = tk.DoubleVar()
        self.trans_type = tk.StringVar(value="Pay")
        self.giver = tk.StringVar()
        self.receiver = tk.StringVar()
        self.explanation = tk.StringVar()
        self.date = tk.StringVar(value=datetime.now().strftime("%d.%m.%Y"))

        self.cur_amount = tk.DoubleVar()
        self.total_rec = tk.DoubleVar()
        self.total_pay = tk.DoubleVar()

        self.style = ttk.Style()
        self.style.theme_use("default")
        # Configure custom tags
        self.style.configure("Treeview", font=("Arial", 13))
        self.style.map("Treeview", background=[("selected", "#D0E4F5")])

        default_font = tkFont.nametofont("TkDefaultFont")
        default_font.configure(size=12)
        # Build GUI
        self.create_widgets()
        self.load_everything()


        self.root.rowconfigure(10, weight=1)
        self.root.columnconfigure(1, weight=1)

    def create_widgets(self):
        # Entry Fields
        ttk.Label(self.root, text="Amount:").grid(row=0, column=0, sticky="nse", padx=10, pady=5)
        ttk.Entry(self.root, textvariable=self.amount).grid(row=0, column=1, sticky="nsw", padx=10, pady=5)
        
        ttk.Label(self.root, text="Type:").grid(row=1, column=0, sticky="nse", padx=10, pady=5)
        trans_type_values = ["Add", "Pay", "Check", "Loan", "Loan Repayment"]
        ttk.Combobox(self.root, textvariable=self.trans_type, values=trans_type_values).grid(row=1, column=1, sticky="nsw", padx=10, pady=5)

        ttk.Label(self.root, text="From:").grid(row=2, column=0, sticky="nse", padx=10, pady=5)
        ttk.Entry(self.root, textvariable=self.giver).grid(row=2, column=1, sticky="nsw", padx=10, pady=5)

        ttk.Label(self.root, text="To:").grid(row=3, column=0, sticky="nse", padx=10, pady=5)
        ttk.Entry(self.root, textvariable=self.receiver).grid(row=3, column=1, sticky="nsw", padx=10, pady=5)

        ttk.Label(self.root, text="Date:").grid(row=4, column=0, sticky="nse", padx=10, pady=5)
        ttk.Entry(self.root, textvariable=self.date).grid(row=4, column=1, sticky="nsw", padx=10, pady=5)

        ttk.Label(self.root, text="Explanation:").grid(row=5, column=0, sticky="nse", padx=10, pady=5)
        ttk.Entry(self.root, textvariable=self.explanation).grid(row=5, column=1, sticky="nsew", padx=10, pady=5)
        

        
        # Buttons
        ttk.Button(self.root, text="Submit", command=self.submit_transaction).grid(row=6, column=0, columnspan=2)
        ttk.Separator(self.root).grid(row=7, column=0, columnspan=2, sticky="nsew", padx=10, pady=10)

        ttk.Button(self.root, text="View All", command=self.load_transactions).grid(row=8, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.root, text="Filter by Date", command=self.filter_by_date).grid(row=9, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.root, text="Filter by Type", command=self.filter_by_type).grid(row=10, column=0, columnspan=2, padx=10, pady=10)
        ttk.Button(self.root, text="Export to PDF", command=self.export_to_docx).grid(row=11, column=0, columnspan=2, padx=10, pady=10)

        # Transaction List Frame with Scrollbar
        tree_frame = ttk.Frame(self.root)
        tree_frame.grid(row=0, column=3, rowspan=14, padx=10, pady=10, sticky="nsew")

        self.tree = ttk.Treeview(tree_frame, columns=("ID", "Amount", "Type", "From", "To", "Explanation", "Date"),
                                 show="headings")
        self.tree.grid(row=0, column=0, sticky="nsew")

        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.tree.configure(yscrollcommand=scrollbar.set)

        # Configure tree headings
        self.tree.heading("ID", text="ID")
        self.tree.heading("Amount", text="Amount")
        self.tree.heading("Type", text="Type")
        self.tree.heading("From", text="From")
        self.tree.heading("To", text="To")
        self.tree.heading("Explanation", text="Explanation")
        self.tree.heading("Date", text="Date")

        self.tree.column("ID", width=50, anchor="center")
        self.tree.column("Amount", width=100, anchor="e")
        self.tree.column("Type", width=100, anchor="w")
        self.tree.column("From", width=150, anchor="w")
        self.tree.column("To", width=150, anchor="w")
        self.tree.column("Explanation", width=200, anchor="w")
        self.tree.column("Date", width=100, anchor="center")

        # Color tags
        self.tree.tag_configure("positive", foreground="green")
        self.tree.tag_configure("negative", foreground="red")

        # Expand behavior
        tree_frame.rowconfigure(0, weight=1)
        tree_frame.columnconfigure(0, weight=1)

        ttk.Button(tree_frame, text="Delete Selected", command=self.delete_selected_from_db).grid(row=1, column=0, columnspan=2, pady=10, padx=10)
        ttk.Separator(tree_frame).grid(row=2, column=0, sticky="nsew", padx=10, pady=10)

        info_frame = ttk.Frame(tree_frame)
        info_frame.grid(row=3, column=0, sticky="nsew", padx=10, pady=10)

        ttk.Label(info_frame, text="Current Amount: ").grid(row=0, column=0, padx=10, pady=10)
        cur_label = ttk.Label(info_frame, textvariable=self.cur_amount)
        cur_label.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(info_frame, text="Total Received: ").grid(row=0, column=2, padx=10, pady=10)
        total_received = ttk.Label(info_frame, textvariable=self.total_rec)
        total_received.grid(row=0, column=3, padx=10, pady=10)

        ttk.Label(info_frame, text="Total Paid: ").grid(row=0, column=4, padx=10, pady=10)
        total_paid = ttk.Label(info_frame, textvariable=self.total_pay)
        total_paid.grid(row=0, column=5, padx=10, pady=10)


        # Totals
        ttk.Label(self.root, text="hey")
    
    def submit_transaction(self):
        try:
            if self.trans_type.get() in ("Pay", "Check", "Loan Repayment"):
                self.amount.set(self.amount.get() * -1)
            add_transaction(
                self.amount.get(),
                self.trans_type.get(),
                self.giver.get(),
                self.receiver.get(),
                self.explanation.get(),
                self.date.get()
            )
            self.load_everything()
            self.amount.set(0)
        except Exception as e:
            messagebox.showerror("Error", str(e))
    
    def load_transactions(self):
        for row in self.tree.get_children():
            self.tree.delete(row)
        for trans in get_all_transactions():
            amount = trans[1]
            tag = "positive" if amount >= 0 else "negative"
            self.tree.insert("", "end", values=trans, tags=tag)

    def load_current_amount(self):
        cur = get_total_amount_transactions_sum()
        self.cur_amount.set(cur)

    def load_total_paid(self):
        cur = get_total_amount_transactions_paid()
        self.total_pay.set(cur)

    def load_total_received(self):
        cur = get_total_transactions_received()
        self.total_rec.set(cur)

    def load_everything(self):
        self.load_transactions()
        self.load_current_amount()
        self.load_total_paid()
        self.load_total_received()
    
    def filter_by_date(self):
        date = self.date.get()
        for row in self.tree.get_children():
            self.tree.delete(row)
        for trans in get_transactions_by_date(date):
            amount = trans[1]
            tag = "positive" if amount >= 0 else "negative"
            self.tree.insert("", "end", values=trans, tags=tag)

    def filter_by_type(self):
        type = self.trans_type.get()
        for row in self.tree.get_children():
            self.tree.delete(row)
        for trans in get_transactions_by_type(type):
            amount = trans[1]
            tag = "positive" if amount >= 0 else "negative"
            self.tree.insert("", "end", values=trans, tags=tag)

    def delete_selected_from_db(self):
        cur_item = self.tree.focus()
        if self.tree.item(cur_item, "values") == '':
            tk.messagebox.showerror("Error", "You must select a transaction")
            return
        cur_id = self.tree.item(cur_item, "values")[0]
        delete_transaction_by_id(cur_id)
        self.load_everything()

    def export_to_docx(self, filename="transactions_report.docx"):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import arabic_reshaper
            from bidi.algorithm import get_display
            from datetime import datetime

            # Get all transaction data
            transactions = get_all_transactions()
            receivers = get_receivers()
            current_amount = get_total_amount_transactions_sum()[0][0] or 0
            total_received = get_total_transactions_received()[0][0] or 0
            total_paid = abs(get_total_amount_transactions_paid()[0][0] or 0)

            # Create a new Document
            doc = Document()

            section = doc.sections[0]
            section.top_margin = Inches(0.5)  # 0.5 inches
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

            # Add title
            title = doc.add_heading('Transaction Report', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add subtitle with date
            subtitle = doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(12)

            # Add separator
            doc.add_paragraph("", style='Intense Quote')

            # Add summary information
            doc.add_paragraph("\nSummary Information:", style='Heading 2')

            # Create a table for summary
            summary_table = doc.add_table(rows=5, cols=2)
            summary_table.style = 'Light Shading Accent 1'

            # Summary table data
            summary_cells = summary_table.rows[0].cells
            summary_cells[0].text = "Current Balance"
            summary_cells[1].text = f"{current_amount:,.2f}"

            summary_cells = summary_table.rows[1].cells
            summary_cells[0].text = "Total Received"
            summary_cells[1].text = f"{total_received:,.2f}"

            summary_cells = summary_table.rows[2].cells
            summary_cells[0].text = "Total Paid"
            summary_cells[1].text = f"{total_paid:,.2f}"

            summary_cells = summary_table.rows[3].cells
            summary_cells[0].text = "Total "

            # Add separator
            doc.add_paragraph("", style='Intense Quote')

            # Add transactions header
            doc.add_paragraph("\nTransaction Details:", style='Heading 2')

            # Create transactions table
            trans_table = doc.add_table(rows=1, cols=7)
            trans_table.style = 'Table Grid'

            # Table headers
            trans_hdr_cells = trans_table.rows[0].cells
            headers = ['ID', 'Amount', 'Type', 'From', 'To', 'Explanation', 'Date']
            for i, header in enumerate(headers):
                trans_hdr_cells[i].text = header
                trans_hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                trans_hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Format Arabic text
            def format_arabic(text):
                if text is None:
                    return ""
                text = str(text)
                if any("\u0600" <= c <= "\u06FF" for c in text):
                    reshaped = arabic_reshaper.reshape(text)
                    return get_display(reshaped)
                return text

            # Add transaction data
            for trans in transactions:
                row_cells = trans_table.add_row().cells
                row_cells[0].text = str(trans[0])  # ID
                row_cells[1].text = f"{trans[1]:,.2f}"  # Amount
                row_cells[2].text = trans[2]  # Type
                row_cells[3].text = format_arabic(trans[3])  # From
                row_cells[4].text = format_arabic(trans[4])  # To
                row_cells[5].text = format_arabic(trans[5])  # Explanation
                row_cells[6].text = trans[6]  # Date

                # Set amount color
                amount_para = row_cells[1].paragraphs[0]
                amount_run = amount_para.runs[0]
                amount_run.font.color.rgb = RGBColor(0, 128, 0) if trans[1] >= 0 else RGBColor(255, 0, 0)

                # Right-align all cells
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Set column widths
            for i, width in enumerate([0.5, 1, 1, 2, 2, 2, 1]):
                for cell in trans_table.columns[i].cells:
                    cell.width = Inches(width)



            for rec in receivers:
                # Add separator
                rec_transactions = get_transactions_by_receiver(rec[0])
                doc.add_page_break()
                rec_total_paid = get_total_received_by_receiver(rec[0])[0][0] * -1
                doc.add_paragraph(rec[0], style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("Total Paid: " + str(rec_total_paid), style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

                rec_trans_table = doc.add_table(rows=1, cols=7)
                rec_trans_table.style = 'Table Grid'

                rec_hdr_cells = rec_trans_table.rows[0].cells
                for i, header in enumerate(headers):
                    rec_hdr_cells[i].text = header
                    rec_hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rec_hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                # Add transaction data
                for trans in rec_transactions:
                    rec_row_cells = rec_trans_table.add_row().cells
                    rec_row_cells[0].text = str(trans[0])  # ID
                    rec_row_cells[1].text = f"{trans[1]:,.2f}"  # Amount
                    rec_row_cells[2].text = trans[2]  # Type
                    rec_row_cells[3].text = format_arabic(trans[3])  # From
                    rec_row_cells[4].text = format_arabic(trans[4])  # To
                    rec_row_cells[5].text = format_arabic(trans[5])  # Explanation
                    rec_row_cells[6].text = trans[6]  # Date
                    # Set amount color
                    rec_amount_para = rec_row_cells[1].paragraphs[0]
                    amount_run = rec_amount_para.runs[0]
                    amount_run.font.color.rgb = RGBColor(0, 128, 0) if trans[1] >= 0 else RGBColor(255, 0, 0)
                    # Right-align all cells
                    for cell in rec_row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    # Set column widths
                    for i, width in enumerate([0.5, 1, 1, 2, 2, 2, 1]):
                        for cell in rec_trans_table.columns[i].cells:
                            cell.width = Inches(width)

            doc.add_page_break()
            doc.add_paragraph("Checks Payment", style='Heading 2').alignment = WD_ALIGN_PARAGRAPH.CENTER
            checks_table = doc.add_table(rows=1, cols=7)
            checks_table.style = 'Table Grid'

            checks_hdr_cells = checks_table.rows[0].cells
            for i, header in enumerate(headers):
                checks_hdr_cells[i].text = header
                checks_hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                checks_hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            for trans in get_transaction_checks():
                check_row_cells = checks_table.add_row().cells
                check_row_cells[0].text = str(trans[0])
                check_row_cells[1].text = f"{trans[1]:,.2f}"
                check_row_cells[2].text = trans[2]
                check_row_cells[3].text = format_arabic(trans[3])
                check_row_cells[4].text = format_arabic(trans[4])
                check_row_cells[5].text = format_arabic(trans[5])
                check_row_cells[6].text = trans[6]
                check_amount_para = check_row_cells[1].paragraphs[0]
                check_amount_runs = check_amount_para.runs[0]
                check_amount_runs.font.color.rgb = RGBColor(0, 128, 0) if trans[1] >= 0 else RGBColor(255, 0, 0)
                for cell in checks_table.columns[i].cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                for i, width in enumerate([0.5, 1, 1, 2, 2, 2, 1]):
                    for cell in checks_table.columns[i].cells:
                        cell.width = Inches(width)

            for giv in get_givers():
                print(get_givers())
                doc.add_page_break()
                giv_transactions = get_transactions_by_giver(giv[0])
                giv_total_added = get_total_sum_by_giver(giv[0])
                doc.add_paragraph(giv[0], style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("Total Added: " + str(giv_total_added[0][0]), style="Heading 2").alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                giv_trans_table = doc.add_table(rows=1, cols=7)
                giv_trans_table.style = 'Table Grid'

                giv_hdr_cells = giv_trans_table.rows[0].cells
                for i, header in enumerate(headers):
                    giv_hdr_cells[i].text = header
                    giv_hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    giv_hdr_cells[i].paragraphs[0].runs[0].font.bold = True

                for trans in giv_transactions:
                    giv_trans_cells = giv_trans_table.add_row().cells
                    giv_trans_cells[0].text = str(trans[0])
                    giv_trans_cells[1].text = f"{trans[1]:,.2f}"
                    giv_trans_cells[2].text = trans[2]
                    giv_trans_cells[3].text = format_arabic(trans[3])
                    giv_trans_cells[4].text = format_arabic(trans[4])
                    giv_trans_cells[5].text = format_arabic(trans[5])
                    giv_trans_cells[6].text = trans[6]

                    giv_amount_para = giv_trans_cells[1].paragraphs[0]
                    amount_run = giv_amount_para.runs[0]
                    amount_run.font.color.rgb = RGBColor(0, 128, 0) if trans[1] >= 0 else RGBColor(255, 0, 0)
                    # Right-align all cells
                    for cell in giv_trans_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    # Set column widths
                    for i, width in enumerate([0.5, 1, 1, 2, 2, 2, 1]):
                        for cell in giv_trans_table.columns[i].cells:
                            cell.width = Inches(width)

            doc.add_page_break()
            add_money_trans = get_transaction_add_money()
            doc.add_paragraph("Added Money", style='Heading 1').alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            add_money_table = doc.add_table(rows=1, cols=7)
            add_money_table.style = 'Table Grid'

            add_money_hdr_cells = add_money_table.rows[0].cells
            for i, header in enumerate(headers):
                add_money_hdr_cells[i].text = header
                add_money_hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_money_hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            for trans in add_money_trans:
                add_money_cells = add_money_table.add_row().cells
                add_money_cells[0].text = str(trans[0])
                add_money_cells[1].text = f"{trans[1]:,.2f}"
                add_money_cells[2].text = trans[2]
                add_money_cells[3].text = format_arabic(trans[3])
                add_money_cells[4].text = format_arabic(trans[4])
                add_money_cells[5].text = format_arabic(trans[5])
                add_money_cells[6].text = trans[6]

                add_money_amount_para = add_money_cells[1].paragraphs[0]
                amount_run = add_money_amount_para.runs[0]
                amount_run.font.color.rgb = RGBColor(0, 128, 0) if trans[1] >= 0 else RGBColor(255, 0, 0)
                # Right-align all cells
                for cell in add_money_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

                for i, width in enumerate([0.5, 1, 1, 2, 2, 2, 1]):
                    for cell in add_money_table.columns[i].cells:
                        cell.width = Inches(width)

            # Save the document
            doc.save(filename)
            messagebox.showinfo("Export Successful", f"DOCX exported successfully as {filename}")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to export DOCX: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = MoneyManagerApp(root)
    root.mainloop()